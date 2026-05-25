import re
from pathlib import Path
from uuid import uuid4

import pdfplumber
from docx import Document
from fastapi import HTTPException, UploadFile, status
from PyPDF2 import PdfReader

from app.core.config import settings


ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_UPLOAD_SIZE = 8 * 1024 * 1024


async def save_upload(file: UploadFile) -> Path:
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF and DOCX resume files are supported.",
        )

    payload = await file.read()
    if len(payload) > MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, detail="File exceeds 8MB limit.")
    if not payload:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty.")

    safe_name = re.sub(r"[^A-Za-z0-9_.-]", "_", file.filename or f"resume{suffix}")
    target = settings.upload_dir / f"{uuid4().hex}_{safe_name}"
    target.write_bytes(payload)
    return target


def extract_text_from_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix == ".docx":
        return _extract_docx_text(path)
    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported file type.")


def _extract_pdf_text(path: Path) -> str:
    try:
        with pdfplumber.open(path) as pdf:
            pages = [page.extract_text(x_tolerance=1, y_tolerance=3) or "" for page in pdf.pages]
        text = "\n".join(pages).strip()
        if text:
            return text
    except Exception:
        pass

    try:
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Could not parse PDF text.") from exc


def _extract_docx_text(path: Path) -> str:
    try:
        document = Document(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_lines = []
        for table in document.tables:
            for row in table.rows:
                table_lines.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        return "\n".join([*paragraphs, *table_lines]).strip()
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Could not parse DOCX text.") from exc
